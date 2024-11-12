import streamlit as st
import numpy as np
import cv2
from PIL import Image
import pytesseract
import re
from sympy import latex, sympify
import docx
import io
import tempfile
import pyperclip
import platform

# Cấu hình Streamlit
st.set_page_config(
    page_title="Math OCR & LaTeX Converter",
    page_icon="📝",
    layout="wide"
)

# CSS tùy chỉnh
st.markdown("""
<style>
    .main { 
        padding: 2rem; 
    }
    .result-box {
        background-color: #f0f2f6;
        border-radius: 10px;
        padding: 20px;
        margin: 10px 0;
    }
    .latex-output {
        background-color: white;
        padding: 15px;
        border-radius: 5px;
        font-family: monospace;
    }
</style>
""", unsafe_allow_html=True)

# Hàm xử lý ảnh
def preprocess_image(image):
    # Chuyển về grayscale
    gray = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2GRAY)
    
    # Tăng cường độ tương phản
    clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
    enhanced = clahe.apply(gray)
    
    # Khử nhiễu
    denoised = cv2.fastNlMeansDenoising(enhanced)
    
    # Ngưỡng hóa
    _, binary = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    
    return binary

# Hàm tách biểu thức toán học
def extract_math(text):
    math_patterns = [
        r'y\s*=\s*\w+[^.]',  # Phương trình
        r'\d+\s[+-/]\s\d+',  # Phép tính cơ bản
        r'\w+\s*=\s*[^.]*',  # Biểu thức gán
        r'\w+\s*[+\-*/]\s*\w+',  # Biểu thức đơn giản
    ]
    
    math_expressions = []
    for pattern in math_patterns:
        matches = re.finditer(pattern, text)
        for match in matches:
            math_expressions.append(match.group())
    
    return math_expressions

# Hàm chuyển đổi sang LaTeX
def math_to_latex(expression):
    try:
        expression = expression.replace('=', ' = ')
        latex_expr = latex(sympify(expression, evaluate=False))
        return f"${latex_expr}$"
    except:
        return f"${expression}$"

# Hàm tạo tài liệu Word
def create_word_doc(original_text, latex_code):
    doc = docx.Document()
    doc.add_heading('Kết quả chuyển đổi toán học', 0)
    
    doc.add_heading('Văn bản gốc:', level=1)
    doc.add_paragraph(original_text)
    
    doc.add_heading('Mã LaTeX:', level=1)
    doc.add_paragraph(latex_code)
    
    return doc

# Hàm chụp màn hình (hỗ trợ đa nền tảng)
def take_screenshot():
    system = platform.system()
    
    if system == "Darwin":  # macOS
        import pyscreenshot as ImageGrab
        image = ImageGrab.grab()
    elif system == "Windows":
        from PIL import ImageGrab
        image = ImageGrab.grab()
    elif system == "Linux":
        import pyscreenshot as ImageGrab
        image = ImageGrab.grab()
    else:
        st.error("Hệ điều hành không được hỗ trợ")
        return None
    
    return image

# Hàm chính
def main():
    st.title("📚 Chuyển đổi Toán học sang LaTeX")
    
    # Các tab chức năng
    tab1, tab2, tab3 = st.tabs(["Tải ảnh", "Chụp màn hình", "Hướng dẫn"])
    
    with tab1:
        # Upload ảnh
        uploaded_file = st.file_uploader("Tải lên ảnh chứa nội dung toán học", 
                                         type=['png', 'jpg', 'jpeg'])
        
        if uploaded_file is not None:
            image = Image.open(uploaded_file)
            st.image(image, caption="Ảnh đã tải lên", use_column_width=True)
            
            if st.button("Xử lý ảnh từ tập tin"):
                process_image(image)
    
    with tab2:
        if st.button("Chụp màn hình"):
            screenshot = take_screenshot()
            if screenshot:
                st.image(screenshot, caption="Ảnh chụp màn hình", use_column_width=True)
                process_image(screenshot)
    
    with tab3:
        st.markdown("""
        ### Hướng dẫn sử dụng
        1. Tải ảnh chứa công thức toán học
        2. Nhấn nút "Xử lý ảnh"
        3. Xem kết quả và tải về định dạng DOCX
        
        ### Lưu ý
        - Hỗ trợ các định dạng: PNG, JPG, JPEG
        - Chất lượng ảnh ảnh hưởng đến độ chính xác
        """)

# Hàm xử lý ảnh chung
def process_image(image):
    with st.spinner("Đang xử lý..."):
        # Tiền xử lý ảnh
        processed_img = preprocess_image(image)
        
        # OCR
        text = pytesseract.image_to_string(processed_img, lang='vie+eng')
        
        # Tách biểu thức toán học
        math_expressions = extract_math(text)
        
        # Chuyển đổi sang LaTeX
        latex_results = [math_to_latex(expr) for expr in math_expressions]
        
        # Hiển thị kết quả
        st.subheader("Kết quả nhận dạng")
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("Văn bản gốc:")
            st.code(text)
        
        with col2:
            st.markdown("Mã LaTeX:")
            latex_text = "\n".join(latex_results)
            st.code(latex_text)
            
            # Nút copy LaTeX
            if st.button("📋 Copy LaTeX"):
                pyperclip.copy(latex_text)
                st.success("Đã copy mã LaTeX!")
        
        # Tạo file Word
        doc = create_word_doc(text, latex_text)
        
        # Lưu tạm file Word
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            doc.save(tmp.name)
        
        # Nút tải xuống
        with open(tmp.name, 'rb') as f:
            st.download_button(
                label="📥 Tải kết quả (DOCX)",
                data=f.read(),
                file_name="math_conversion.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# Chạy ứng dụng
if __name__ == "__main__":
    main()
